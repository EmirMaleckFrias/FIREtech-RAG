"""Troceo con conciencia de párrafo: PDF por párrafos reales y tablas de Word
que no cambian los números de columna.

Dos defectos reproducidos con el parser real el 4 sep 2026:

1. El PDF se troceaba por LÍNEAS físicas, sin solape y sin arrastrar el
   encabezado. "between 2018 and" | "2021." quedaba partido, "hippocam-" |
   "pal" hacía que "hippocampal" no existiera en el índice, y la frontera de
   ~400 tokens caía en una línea arbitraria (podía separar "542" de "pg/mL").

2. Las tablas de Word deduplicaban celdas adyacentes por IGUALDAD DE TEXTO
   creyendo que eran celdas combinadas. Con dos grupos de la misma edad media,
   "Edad | 72.4 (6.1) | 72.4 (6.1) | 74.0 (5.8) | 0.31" salía con una columna
   menos y 74.0 pasaba a leerse bajo el grupo equivocado.

Y cuatro más, medidos el 4 sep 2026 sobre el propio parser (los dos primeros
eran REGRESIONES de la primera versión de la reconstrucción de párrafos, medidas
contra el troceo por líneas que había antes):

3. Las tablas dentro de un PDF perdían las filas: al no acabar en punto se
   pegaban unas a otras y salía una sola línea con todos los números seguidos,
   ninguno junto a su cabecera.
4. Las celdas combinadas de verdad (gridSpan) colapsaban a UNA posición, así
   que la fila tenía menos columnas que la cabecera y el último valor se leía
   bajo la cabecera equivocada.
5. La cabecera que se repite en cada bloque de una tabla larga era ciegamente
   la fila 0, que en Word suele ser un título combinado a todo el ancho: los
   bloques 2..N repetían el título y no llevaban los nombres de columna.
6. El rótulo de tabla se heredaba de una tabla a la siguiente, el punto de
   "Table 1." partía una oración que solo la citaba, la cita Vancouver pegada
   al punto ("decline.12,13") impedía cerrar la oración, y un compuesto
   partido en su propio guion perdía el guion ("antiinflammatory").

Se generan archivos de verdad (PDF a mano, DOCX con python-docx) y se parsean
por el camino real, que es la única forma de comprobar que las heurísticas
sobreviven a un archivo.
"""
from __future__ import annotations

from pathlib import Path

import pytest

from app.ingest.generic import (
    _OVERLAP_TOKENS,
    _abre_parrafo,
    _cierra_oracion,
    _est_tokens,
    _parece_fila_de_tabla,
    _unir_lineas,
    parse_generic,
)
from tests.pdf_falso import escribir_pdf

TITULO = "Cerebrospinal fluid biomarkers in early Alzheimer disease"

PORTADA = [
    (TITULO, 17.0),
    ("Ricardo F. Allegri, Manuel Colome, Juan C. Guilbe", 11.0),
    ("doi:10.3233/JAD-220123  J Alzheimers Dis 2023", 8.0),
]


def _texto(chunks: list[dict]) -> str:
    return "\n".join(c["text"] for c in chunks)


def _cuerpo(chunk: dict) -> str:
    """Texto del chunk sin las líneas de contexto (título y sección)."""
    return chunk["text"].split("\n\n", 1)[1] if "\n\n" in chunk["text"] else chunk["text"]


def _de_seccion(chunks: list[dict], seccion: str) -> list[dict]:
    return [c for c in chunks if c["section"] == seccion]


def _solape(anterior: dict, siguiente: dict) -> list[str]:
    """Párrafos finales de `anterior` con los que arranca `siguiente`.

    El solape son ~_OVERLAP_TOKENS, o sea VARIOS párrafos cortos de cola, no
    solo el último: se busca el sufijo más largo del anterior que sea prefijo
    del siguiente.
    """
    ant = _cuerpo(anterior).split("\n\n")
    sig = _cuerpo(siguiente).split("\n\n")
    for k in range(min(len(ant), len(sig)), 0, -1):
        if ant[-k:] == sig[:k]:
            return sig[:k]
    return []


# ---------------------------------------------------------------------------
# PDF: párrafos reconstruidos a partir de líneas físicas
# ---------------------------------------------------------------------------
def test_las_lineas_fisicas_se_unen_en_parrafos(tmp_path: Path):
    """La reproducción literal del defecto: oración partida y palabra cortada."""
    pdf = escribir_pdf(tmp_path / "metodos.pdf", [[
        *PORTADA,
        ("Methods", 12.0),
        ("We recruited 120 participants aged between 55 and 85 years between 2018 and", 10.0),
        ("2021. Volumetric MRI focused on the hippocam-", 10.0),
        ("pal formation and adjacent cortex.", 10.0),
    ]])

    chunks, _ = parse_generic(pdf, pdf.name)
    metodos = _de_seccion(chunks, "Methods")

    assert len(metodos) == 1
    texto = metodos[0]["text"]
    assert "between 2018 and 2021." in texto
    assert "hippocampal formation" in texto
    assert "hippocam-" not in texto
    # Y no se pegan con doble salto lo que era una sola oración.
    assert "and\n\n2021" not in texto


def test_una_cifra_no_se_separa_de_su_unidad(tmp_path: Path):
    pdf = escribir_pdf(tmp_path / "resultados.pdf", [[
        *PORTADA,
        ("Results", 12.0),
        ("Mean amyloid beta 42 was 542", 10.0),
        ("pg/mL in the impaired group and 912 pg/mL in controls.", 10.0),
    ]])

    chunks, _ = parse_generic(pdf, pdf.name)

    assert "542 pg/mL in the impaired group" in _texto(chunks)


def test_un_parrafo_que_cruza_de_pagina_conserva_las_dos_paginas(tmp_path: Path):
    """La oración sigue en la página siguiente: el chunk cita ambas."""
    pdf = escribir_pdf(tmp_path / "salto.pdf", [
        [
            *PORTADA,
            ("Methods", 12.0),
            ("Participants were enrolled at three memory clinics between 2018 and", 10.0),
        ],
        [
            ("2021. Volumetric MRI focused on the hippocam-", 10.0),
            ("pal formation and adjacent cortex.", 10.0),
        ],
    ])

    chunks, _ = parse_generic(pdf, pdf.name)
    metodos = _de_seccion(chunks, "Methods")

    assert len(metodos) == 1
    assert "between 2018 and 2021." in metodos[0]["text"]
    assert "hippocampal" in metodos[0]["text"]
    assert metodos[0]["page"] == 1
    assert metodos[0]["source_pages"] == [1, 2]


def test_el_guion_se_conserva_cuando_forma_parte_del_termino(tmp_path: Path):
    """"anti-" + "Alzheimer" y "COVID-" + "19" no son palabras cortadas."""
    pdf = escribir_pdf(tmp_path / "guiones.pdf", [[
        *PORTADA,
        ("Methods", 12.0),
        ("Prior treatment with anti-", 10.0),
        ("Alzheimer agents and any COVID-", 10.0),
        ("19 vaccination were recorded at baseline.", 10.0),
    ]])

    chunks, _ = parse_generic(pdf, pdf.name)
    texto = _texto(chunks)

    assert "anti-Alzheimer agents" in texto
    assert "COVID-19 vaccination" in texto


def test_cada_chunk_lleva_titulo_y_seccion_dentro_del_texto(tmp_path: Path):
    """El contexto viaja en lo que se embebe, no solo en el payload."""
    pdf = escribir_pdf(tmp_path / "contexto.pdf", [[
        *PORTADA,
        ("Methods", 12.0),
        ("We recruited 120 participants from three memory clinics.", 10.0),
        ("Results", 12.0),
        ("Mean amyloid beta 42 was 542 pg/mL in the impaired group.", 10.0),
        ("Discussion", 12.0),
        ("These findings may indicate an earlier onset than assumed.", 10.0),
    ]])

    chunks, _ = parse_generic(pdf, pdf.name)

    for seccion in ("Methods", "Results", "Discussion"):
        (chunk,) = _de_seccion(chunks, seccion)
        assert chunk["text"].startswith(f"{TITULO}\n{seccion}\n\n"), chunk["text"]
    # El contexto no contamina la sección del payload ni la cita.
    assert all(c["title"] == TITULO for c in chunks)
    assert all(c["citation"] == "Allegri et al., 2023" for c in chunks)


def test_el_titulo_no_se_repite_cuando_ya_es_la_seccion_vigente(tmp_path: Path):
    """En la portada el título se detecta como encabezado por formato y pasa a
    ser la sección: no debe salir dos veces seguidas."""
    pdf = escribir_pdf(tmp_path / "portada.pdf", [[
        ("Tau load and cognition", 16.0),
        ("Higher tau load was associated with faster cognitive decline.", 10.0),
    ]])

    chunks, _ = parse_generic(pdf, pdf.name)

    assert len(chunks) == 1
    assert chunks[0]["section"] == "Tau load and cognition"
    assert chunks[0]["text"].count("Tau load and cognition") == 1
    assert chunks[0]["text"].startswith("Tau load and cognition\n\nHigher tau")


def test_sin_titulo_ni_seccion_el_texto_queda_limpio(tmp_path: Path):
    pdf = escribir_pdf(tmp_path / "plano.pdf", [[
        ("Primera linea del documento sin ninguna estructura visible.", 10.0),
        ("Segunda linea que continua el mismo parrafo de siempre.", 10.0),
    ]])

    chunks, _ = parse_generic(pdf, pdf.name)

    assert len(chunks) == 1
    assert chunks[0]["text"].startswith("Primera linea")
    assert not chunks[0]["text"].startswith("\n")


def _oraciones(prefijo: str, n: int) -> list[tuple[str, float]]:
    """n líneas de ~80 caracteres, cada una una oración completa."""
    return [
        (f"{prefijo} sentence {i:03d} describes one recruitment step of the cohort in detail.", 10.0)
        for i in range(n)
    ]


def test_los_chunks_se_solapan_dentro_de_una_seccion_y_no_entre_secciones(tmp_path: Path):
    """El solape del 15% prometido por el módulo existe para PDF.

    Y no cruza secciones: el final de Métodos no puede aparecer "de cola" en el
    primer chunk de Resultados, donde se leería como evidencia.
    """
    metodos = _oraciones("Methods", 44)
    pdf = escribir_pdf(tmp_path / "largo.pdf", [
        [*PORTADA, ("Methods", 12.0), *metodos[:22]],
        [*metodos[22:], ("Results", 12.0), ("Mean amyloid beta 42 was 542 pg/mL.", 10.0)],
    ])

    chunks, _ = parse_generic(pdf, pdf.name)
    de_metodos = _de_seccion(chunks, "Methods")

    assert len(de_metodos) >= 2, "44 oraciones de ~20 tokens deben dar mas de un chunk"
    for anterior, siguiente in zip(de_metodos, de_metodos[1:]):
        cola = _solape(anterior, siguiente)
        assert cola, "el chunk siguiente debe arrancar con la cola del anterior"
        # Y esa cola es del orden del 15% prometido (el empaquetador toma
        # párrafos de cola mientras quepan en ~_OVERLAP_TOKENS): ni una línea
        # suelta ni el chunk entero.
        assert _OVERLAP_TOKENS // 2 <= sum(_est_tokens(p) for p in cola) <= 2 * _OVERLAP_TOKENS
        assert len(cola) < len(_cuerpo(anterior).split("\n\n"))
    # Nada de Métodos se cuela en Resultados.
    (resultados,) = _de_seccion(chunks, "Results")
    assert "recruitment step" not in resultados["text"]
    # Y el conjunto sigue conteniendo todas las oraciones (nada se pierde).
    texto = _texto(chunks)
    assert all(f"sentence {i:03d}" in texto for i in range(44))


# ---------------------------------------------------------------------------
# Reglas de unión de líneas, una a una
# ---------------------------------------------------------------------------
@pytest.mark.parametrize(
    "anterior, linea, esperado",
    [
        ("hippocam-", "pal formation", "hippocampal formation"),
        ("anti-", "Alzheimer agents", "anti-Alzheimer agents"),
        # Compuesto partido EN su propio guion: el guion es del término y sin
        # él el término no existe en el índice ("antiinflammatory").
        ("anti-", "inflammatory drugs were allowed", "anti-inflammatory drugs were allowed"),
        ("beta-", "amyloid plaques", "beta-amyloid plaques"),
        ("non-", "carriers of APOE", "non-carriers of APOE"),
        # Izquierda corta: se asume compuesto, no palabra partida.
        ("long-", "term follow up", "long-term follow up"),
        ("COVID-", "19 vaccination", "COVID-19 vaccination"),
        ("doses of 12-", "15 mg", "doses of 12-15 mg"),
        ("between 2018 and", "2021.", "between 2018 and 2021."),
    ],
)
def test_unir_lineas(anterior: str, linea: str, esperado: str):
    assert _unir_lineas(anterior, linea) == esperado


@pytest.mark.parametrize(
    "anterior, linea, abre",
    [
        # Oración cerrada + mayúscula: párrafo nuevo (aunque a veces sea la
        # misma, cortar entre oraciones no hace daño).
        ("decline over time.", "The second cohort was smaller.", True),
        ("was significant (p < 0.05).", "Table 1 shows the baseline data.", True),
        ('as "probable AD."', "We then compared both groups.", True),
        ("decline over time.", "2021 was the last year of follow up.", True),
        # Sin cerrar la oración, nada abre párrafo: ni una cifra ni mayúscula.
        ("between 2018 and", "2021. Volumetric MRI was used.", False),
        ("Mean amyloid beta 42 was 542", "pg/mL in the impaired group.", False),
        ("hippocam-", "pal formation", False),
        ("were assessed by", "Smith and colleagues in 2019.", False),
        # Abreviaturas con punto que no cierran la oración.
        ("as reported by Smith et al.", "(2019) in a larger cohort.", False),
        ("several biomarkers, e.g.", "amyloid and tau.", False),
        # Minúscula tras punto: abreviatura no listada o punto decimal.
        ("as shown in Fig.", "2 of the supplement.", False),
        ("the threshold was approx.", "2.4 pg/mL in both groups.", False),
        # Viñetas y rótulos abren párrafo aunque la anterior no cierre.
        ("inclusion criteria were", "• age over 65 years", True),
        ("Age, years 72.4 (6.1) 74.0 (5.8)", "Table 2. Outcomes at 24 months", True),
        ("Age, years 72.4 (6.1) 74.0 (5.8)", "Figure 3: Kaplan Meier curves", True),
        ("Age, years 72.4 (6.1) 74.0 (5.8)", "TABLE S1 | Sensitivity analyses", True),
        # Una referencia a la tabla DENTRO de la oración no es un rótulo: partir
        # ahí sería volver a separar "in" de "Table 1".
        ("as summarized in", "Table 1, the groups differed at baseline.", False),
        ("differences are shown in", "Table 2 and Figure 3 for both cohorts.", False),
        ("differences are shown in", "Figure 3). No other effect was seen.", False),
        # Un número negativo no es una viñeta.
        ("the coefficient was", "-0.31 (p = 0.02) in the adjusted model.", False),
        # El punto de "Table 1." puede ser el que cierra la FRASE que la cita:
        # si la línea anterior sigue a mitad de oración, no es un rótulo.
        ("summarized in", "Table 1. The groups did not differ.", False),
        ("as reported in", "Figure 2. Curves diverged after 12 months.", False),
        # Cita Vancouver pegada al punto: la oración sí cerró.
        ("associated with faster cognitive decline.12,13", "The second cohort was smaller.", True),
        ("only the first one.14", "A sensitivity analysis followed.", True),
        ("reported before.12-14", "We repeated the model without them.", True),
        # Filas de tabla: no se pegan ni entre ellas ni con lo que las rodea.
        ("Table 1. Baseline characteristics", "Variable Control MCI AD p", True),
        ("Variable Control MCI AD p", "Age, years 72.4 (6.1) 72.4 (6.1) 74.0 (5.8) 0.31", True),
        ("Age, years 72.4 (6.1) 74.0 (5.8) 0.31", "MMSE 28.1 (1.2) 21.3 (3.4) <0.001", True),
        ("MMSE 28.1 (1.2) 21.3 (3.4) <0.001", "Values are mean (SD) unless stated.", True),
        # Subtítulo no detectado: línea corta, sin puntuación, seguida de
        # mayúscula. No debe comerse la primera frase del párrafo.
        ("Statistical analysis", "Group differences were tested with ANOVA.", True),
        ("Sample size calculation", "We assumed a drop out rate of 15%.", True),
        # Y la prosa cortada que también es corta NO es un subtítulo: acaba en
        # palabra de enlace, o la línea siguiente empieza en minúscula.
        ("Cognition was measured with", "Mini Mental State Examination scores.", False),
        ("The risk of dementia was", "higher in carriers of the allele.", False),
        # Ni una primera línea corta que acaba en cifra: es prosa de Resultados,
        # y además el decimal no puede leerse como fin de oración.
        ("The mean difference was 0.31", "Compared with controls it was small.", False),
        ("We used gpt-5.4", "The model was not changed afterwards.", False),
    ],
)
def test_abre_parrafo(anterior: str, linea: str, abre: bool):
    assert _abre_parrafo(anterior, linea) is abre


@pytest.mark.parametrize(
    "linea, es_fila",
    [
        # Tras la etiqueta, todo cifras: fila de tabla.
        ("Age, years 72.4 (6.1) 72.4 (6.1) 74.0 (5.8) 0.31", True),
        ("MMSE 28.1 (1.2) 26.0 (2.1) 21.3 (3.4) <0.001", True),
        ("Female, n (%) 45 (60) 40 (55) 38 (52) 0.71", True),
        ("Amyloid beta 42, pg/mL 912 (210) 640 (180) 542 (150) <0.001", True),
        # Rejilla de espacios (lo que da un PDF maquetado de verdad).
        ("Variable   Control   MCI   AD", True),
        # Prosa con datos: SIEMPRE lleva palabras de enlace entre las cifras.
        ("Mean amyloid beta 42 was 542", False),
        ("We recruited 120 participants aged between 55 and 85 years between 2018 and", False),
        ("the coefficient was -0.31 (p = 0.02) in the adjusted model.", False),
        ("Participants were enrolled between 2018 and 2021 at three centres.", False),
        ("Values are mean (SD) unless stated.", False),
        ("Table 1. Baseline characteristics", False),
    ],
)
def test_parece_fila_de_tabla(linea: str, es_fila: bool):
    assert _parece_fila_de_tabla(linea) is es_fila


@pytest.mark.parametrize(
    "texto, cierra",
    [
        ("decline over time.", True),
        ('as "probable AD."', True),
        # Cita Vancouver pegada al punto: el estilo de casi toda revista médica.
        ("associated with faster cognitive decline.12,13", True),
        ("the first one.14", True),
        ("reported before.12-14", True),
        # Y el decimal y la versión NO cierran oración: llevan un dígito, no una
        # letra, delante del punto. Si cerraran, el párrafo se partiría en
        # mitad de una cifra.
        ("the mean difference was 0.31", False),
        ("we used gpt-5.4", False),
        ("between 2018 and", False),
    ],
)
def test_cierra_oracion(texto: str, cierra: bool):
    assert _cierra_oracion(texto) is cierra


# ---------------------------------------------------------------------------
# PDF: una tabla dentro del PDF conserva sus filas
# ---------------------------------------------------------------------------
def test_una_tabla_de_pdf_conserva_una_fila_por_parrafo(tmp_path: Path):
    """Las dos cosas a la vez, que es donde estaba la regresión.

    Reconstruir párrafos con la regla "si la línea anterior no cierra oración,
    continúa" pegaba TODAS las filas de una tabla en una sola línea (las filas
    no acaban en punto), y con ellas el rótulo y la nota al pie: medido el 4 sep
    2026 contra el troceo por líneas anterior, donde cada fila era su propio
    párrafo. En un RAG médico la tabla de basales es justo donde están los
    datos, así que la fila tiene que seguir siendo una unidad legible.

    Y el párrafo de prosa que la presenta, partido en dos líneas por el
    maquetador y con una referencia "Table 1." que cae a inicio de línea, tiene
    que seguir reconstruyéndose entero.
    """
    pdf = escribir_pdf(tmp_path / "tabla.pdf", [[
        *PORTADA,
        ("Results", 12.0),
        ("Baseline characteristics of the 120 participants are summarized in", 10.0),
        ("Table 1. The three groups did not differ in age or sex distribution.", 10.0),
        ("Table 1. Baseline characteristics", 10.0),
        ("Variable Control MCI AD p", 10.0),
        ("Age, years 72.4 (6.1) 72.4 (6.1) 74.0 (5.8) 0.31", 10.0),
        ("MMSE 28.1 (1.2) 26.0 (2.1) 21.3 (3.4) <0.001", 10.0),
        ("Amyloid beta 42, pg/mL 912 (210) 640 (180) 542 (150) <0.001", 10.0),
        ("Values are mean (SD) unless stated.", 10.0),
    ]])

    chunks, _ = parse_generic(pdf, pdf.name)
    (resultados,) = _de_seccion(chunks, "Results")
    parrafos = _cuerpo(resultados).split("\n\n")

    # 1) Una fila por párrafo, ninguna pegada a la de al lado.
    for fila in (
        "Table 1. Baseline characteristics",
        "Variable Control MCI AD p",
        "Age, years 72.4 (6.1) 72.4 (6.1) 74.0 (5.8) 0.31",
        "MMSE 28.1 (1.2) 26.0 (2.1) 21.3 (3.4) <0.001",
        "Amyloid beta 42, pg/mL 912 (210) 640 (180) 542 (150) <0.001",
        "Values are mean (SD) unless stated.",
    ):
        assert fila in parrafos, parrafos

    # 2) Y la prosa partida en dos líneas sigue siendo un párrafo entero: la
    # referencia "Table 1." de mitad de frase no la corta.
    assert (
        "Baseline characteristics of the 120 participants are summarized in "
        "Table 1. The three groups did not differ in age or sex distribution."
    ) in parrafos


def test_la_prosa_con_muchas_cifras_no_se_toma_por_una_tabla(tmp_path: Path):
    """El adversarial de la regla anterior: un párrafo de Resultados está lleno
    de cifras y NO puede partirse en fragmentos por parecer una tabla, ni
    separar una cifra de su unidad."""
    pdf = escribir_pdf(tmp_path / "cifras.pdf", [[
        *PORTADA,
        ("Results", 12.0),
        ("We recruited 120 participants aged between 55 and 85 years between 2018 and", 10.0),
        ("2021. The mean amyloid beta 42 concentration was 542", 10.0),
        ("pg/mL in the impaired group and 912 pg/mL in the 45 controls.", 10.0),
    ]])

    chunks, _ = parse_generic(pdf, pdf.name)
    (resultados,) = _de_seccion(chunks, "Results")

    assert _cuerpo(resultados) == (
        "We recruited 120 participants aged between 55 and 85 years between "
        "2018 and 2021. The mean amyloid beta 42 concentration was 542 pg/mL "
        "in the impaired group and 912 pg/mL in the 45 controls."
    )


def test_un_subtitulo_no_canonico_no_se_come_la_primera_frase(tmp_path: Path):
    """"Statistical analysis" no lo caza ni `detectar_seccion` (no es un nombre
    canónico) ni el formato (mismo cuerpo de letra que el texto), y al no acabar
    en puntuación se pegaba a la frase siguiente, que es la que mejor lo
    resume."""
    pdf = escribir_pdf(tmp_path / "subtitulo.pdf", [[
        *PORTADA,
        ("Methods", 12.0),
        ("Participants were recruited from three memory clinics between 2018 and 2021.", 10.0),
        ("Statistical analysis", 10.0),
        ("Group differences were tested with analysis of variance and chi square.", 10.0),
    ]])

    chunks, _ = parse_generic(pdf, pdf.name)
    (metodos,) = _de_seccion(chunks, "Methods")
    parrafos = _cuerpo(metodos).split("\n\n")

    assert "Statistical analysis" in parrafos
    assert "Group differences were tested with analysis of variance and chi square." in parrafos


def test_una_oracion_con_cita_vancouver_cierra_el_parrafo(tmp_path: Path):
    """El superíndice de cita se extrae pegado al punto ("decline.12,13"): sin
    reconocerlo como fin de oración el párrafo no cerraba nunca y crecía hasta
    el corte duro de _MAX_PARA_TOKENS, en mitad de una frase."""
    pdf = escribir_pdf(tmp_path / "vancouver.pdf", [[
        *PORTADA,
        ("Discussion", 12.0),
        ("Higher tau load was associated with faster cognitive decline.12,13", 10.0),
        ("The association was independent of amyloid burden.14", 10.0),
        ("Replication in an independent cohort is still needed.", 10.0),
    ]])

    chunks, _ = parse_generic(pdf, pdf.name)
    (discusion,) = _de_seccion(chunks, "Discussion")
    parrafos = _cuerpo(discusion).split("\n\n")

    assert len(parrafos) == 3, parrafos
    assert parrafos[0].endswith("cognitive decline.12,13")


def test_el_guion_del_compuesto_sobrevive_al_corte_de_linea(tmp_path: Path):
    """"anti-" + "inflammatory" quedaba indexado como "antiinflammatory", que no
    lo busca nadie."""
    pdf = escribir_pdf(tmp_path / "compuestos.pdf", [[
        *PORTADA,
        ("Methods", 12.0),
        ("Chronic use of anti-", 10.0),
        ("inflammatory drugs and beta-", 10.0),
        ("amyloid targeting agents was recorded in the hippocam-", 10.0),
        ("pal subgroup.", 10.0),
    ]])

    chunks, _ = parse_generic(pdf, pdf.name)
    texto = _texto(chunks)

    assert "anti-inflammatory drugs" in texto
    assert "beta-amyloid targeting" in texto
    # Y la palabra partida de verdad sigue uniéndose (el compromiso documentado).
    assert "hippocampal subgroup" in texto


# ---------------------------------------------------------------------------
# Word: tablas que no cambian los números de columna
# ---------------------------------------------------------------------------
def _docx():
    return pytest.importorskip("docx")


def test_dos_celdas_iguales_no_combinadas_se_conservan(tmp_path: Path):
    """El caso adversarial: dos grupos con la misma edad media.

    Antes salía "Edad | 72.4 (6.1) | 74.0 (5.8) | 0.31": una columna menos y
    74.0 leído bajo MCI en vez de AD.
    """
    docx = _docx()
    doc = docx.Document()
    doc.add_heading("Results", level=1)
    tabla = doc.add_table(rows=2, cols=5)
    for j, valor in enumerate(["Variable", "Control", "MCI", "AD", "p"]):
        tabla.cell(0, j).text = valor
    for j, valor in enumerate(["Edad", "72.4 (6.1)", "72.4 (6.1)", "74.0 (5.8)", "0.31"]):
        tabla.cell(1, j).text = valor
    destino = tmp_path / "basales.docx"
    doc.save(destino)

    chunks, _ = parse_generic(destino, destino.name)
    (tabla_chunk,) = [c for c in chunks if c["chunk_type"] == "table"]
    filas = _cuerpo(tabla_chunk).splitlines()

    assert filas[0] == "Variable | Control | MCI | AD | p"
    assert filas[1] == "Edad | 72.4 (6.1) | 72.4 (6.1) | 74.0 (5.8) | 0.31"
    # Misma cantidad de columnas en cabecera y datos: nada se desplaza.
    assert filas[0].count("|") == filas[1].count("|")


def test_una_celda_combinada_de_verdad_ocupa_todas_sus_columnas(tmp_path: Path):
    """La combinada horizontal ocupa sus k columnas: texto en la primera y
    marcador vacío en las demás.

    Este test esperaba antes "Control | 72.4 | 72.4" bajo una cabecera de
    CUATRO columnas ("Grupo | Medida | Basal | Final"), o sea fijaba el defecto
    que decía arreglar: colapsar la combinada a una sola posición deja la fila
    con tres columnas y el "72.4" de la columna Final se lee como Basal, que es
    la misma clase de error (un número bajo la cabecera equivocada) que motivó
    dejar de deduplicar por texto. Lo que se comprueba ahora es que las
    posiciones cuadran: el valor de Final sigue siendo el cuarto.

    La combinada vertical (vMerge) sí repite su texto en cada fila, a
    propósito, para que cada fila se lea sola.
    """
    docx = _docx()
    doc = docx.Document()
    tabla = doc.add_table(rows=3, cols=4)
    for j, valor in enumerate(["Grupo", "Medida", "Basal", "Final"]):
        tabla.cell(0, j).text = valor
    horizontal = tabla.cell(1, 1).merge(tabla.cell(1, 2))
    horizontal.text = "72.4"
    tabla.cell(1, 0).text = "Control"
    tabla.cell(1, 3).text = "72.4"
    vertical = tabla.cell(1, 0).merge(tabla.cell(2, 0))
    vertical.text = "Control"
    tabla.cell(2, 1).text = "MMSE"
    tabla.cell(2, 2).text = "28"
    tabla.cell(2, 3).text = "28"
    destino = tmp_path / "combinadas.docx"
    doc.save(destino)

    chunks, _ = parse_generic(destino, destino.name)
    (tabla_chunk,) = [c for c in chunks if c["chunk_type"] == "table"]
    filas = _cuerpo(tabla_chunk).splitlines()

    assert filas[0] == "Grupo | Medida | Basal | Final"
    # La combinada horizontal lleva el texto una vez y deja libre la columna
    # que también abarca; el "72.4" de Final, que tiene el mismo texto pero es
    # otra celda, se conserva Y sigue siendo el cuarto valor.
    assert filas[1] == "Control | 72.4 |  | 72.4"
    assert filas[1].split(" | ")[3] == "72.4"
    # La combinada vertical repite el rótulo de grupo en su segunda fila.
    assert filas[2] == "Control | MMSE | 28 | 28"
    # Y ninguna fila tiene menos columnas que la cabecera.
    assert all(len(f.split(" | ")) == 4 for f in filas)


def test_una_combinada_en_medio_de_la_fila_no_desplaza_la_ultima_columna(tmp_path: Path):
    """Una celda que abarca dos columnas ocupa DOS posiciones.

    Medido el 4 sep 2026: con la cabecera "Grupo | Basal | Final | p" y la
    celda de Basal+Final combinada, la fila salía como "AD | n=40 (both
    visits) | 0.03" (tres columnas contra cuatro) y el p-valor 0.03 se leía
    como el valor de "Final". El texto va en la primera columna que abarca y
    las demás quedan vacías: repetirlo pondría un valor bajo una cabecera en la
    que no se midió.
    """
    docx = _docx()
    doc = docx.Document()
    doc.add_heading("Results", level=1)
    tabla = doc.add_table(rows=2, cols=4)
    for j, valor in enumerate(["Grupo", "Basal", "Final", "p"]):
        tabla.cell(0, j).text = valor
    tabla.cell(1, 0).text = "AD"
    tabla.cell(1, 1).merge(tabla.cell(1, 2)).text = "n=40 (both visits)"
    tabla.cell(1, 3).text = "0.03"
    destino = tmp_path / "combinada_media.docx"
    doc.save(destino)

    chunks, _ = parse_generic(destino, destino.name)
    (tabla_chunk,) = [c for c in chunks if c["chunk_type"] == "table"]
    filas = _cuerpo(tabla_chunk).splitlines()

    assert filas[0].split(" | ") == ["Grupo", "Basal", "Final", "p"]
    assert filas[1].split(" | ") == ["AD", "n=40 (both visits)", "", "0.03"]
    # El p-valor sigue bajo "p", que es lo único que importa de verdad.
    assert filas[1].split(" | ")[3] == "0.03"


def test_una_combinada_al_final_de_la_fila_no_mueve_los_valores(tmp_path: Path):
    """Adversarial de la regla anterior: la combinada llega hasta la última
    columna, así que las posiciones que rellena se recortan al montar la línea
    (no llevan ningún valor). Lo que no puede pasar es que se mueva un valor de
    los de antes."""
    docx = _docx()
    doc = docx.Document()
    tabla = doc.add_table(rows=2, cols=4)
    for j, valor in enumerate(["Variable", "Control", "AD", "p"]):
        tabla.cell(0, j).text = valor
    tabla.cell(1, 0).text = "MMSE"
    tabla.cell(1, 1).text = "28.1"
    tabla.cell(1, 2).merge(tabla.cell(1, 3)).text = "not tested"
    destino = tmp_path / "combinada_final.docx"
    doc.save(destino)

    chunks, _ = parse_generic(destino, destino.name)
    (tabla_chunk,) = [c for c in chunks if c["chunk_type"] == "table"]
    filas = _cuerpo(tabla_chunk).splitlines()

    celdas = filas[1].split(" | ")
    assert celdas[:3] == ["MMSE", "28.1", "not tested"]
    # "28.1" sigue siendo el valor de Control, no de AD.
    assert filas[0].split(" | ")[1] == "Control" and celdas[1] == "28.1"


def test_una_fila_que_empieza_en_la_segunda_columna_no_desplaza_sus_valores(tmp_path: Path):
    """Word permite que una fila empiece más allá de la primera columna
    (`w:gridBefore`), y python-docx no devuelve esos huecos en `fila.cells`.

    Sin reponerlos, la fila salía con una celda menos y el 28 (que es de
    Control) se leía bajo Variable: la misma clase de error que la celda
    combinada colapsada, por la misma razón (contar celdas en vez de columnas
    de la cuadrícula).
    """
    docx = _docx()
    from docx.oxml.ns import qn

    doc = docx.Document()
    tabla = doc.add_table(rows=3, cols=3)
    for j, valor in enumerate(["Variable", "Control", "AD"]):
        tabla.cell(0, j).text = valor
    for j, valor in enumerate(["Edad", "72", "74"]):
        tabla.cell(1, j).text = valor
    for j, valor in enumerate(["MMSE", "28", "21"]):
        tabla.cell(2, j).text = valor
    # La última fila arranca en la segunda columna: se le quita la primera
    # celda y se declara el hueco de rejilla, que es como lo guarda Word.
    tr = tabla.rows[2]._tr
    tr.remove(tr.tc_lst[0])
    trPr = tr.get_or_add_trPr()
    trPr.append(trPr.makeelement(qn("w:gridBefore"), {qn("w:val"): "1"}))
    destino = tmp_path / "hueco.docx"
    doc.save(destino)

    if not hasattr(docx.Document(str(destino)).tables[0].rows[0], "grid_cols_before"):
        pytest.skip("python-docx sin grid_cols_before")

    chunks, _ = parse_generic(destino, destino.name)
    (tabla_chunk,) = [c for c in chunks if c["chunk_type"] == "table"]
    filas = _cuerpo(tabla_chunk).splitlines()

    assert filas[2].split(" | ") == ["", "28", "21"]
    # El 28 sigue bajo Control, que es la columna en la que se midió.
    assert filas[0].split(" | ")[1] == "Control"


def test_una_celda_vacia_al_inicio_de_la_cabecera_no_desplaza_las_columnas(tmp_path: Path):
    """La esquina superior izquierda vacía es una columna, no un adorno."""
    docx = _docx()
    doc = docx.Document()
    tabla = doc.add_table(rows=2, cols=3)
    tabla.cell(0, 1).text = "Control"
    tabla.cell(0, 2).text = "MCI"
    for j, valor in enumerate(["Edad", "72", "73"]):
        tabla.cell(1, j).text = valor
    destino = tmp_path / "esquina.docx"
    doc.save(destino)

    chunks, _ = parse_generic(destino, destino.name)
    (tabla_chunk,) = [c for c in chunks if c["chunk_type"] == "table"]
    filas = _cuerpo(tabla_chunk).splitlines()

    assert filas[0].count("|") == filas[1].count("|") == 2
    assert filas[0].split(" | ")[1:] == ["Control", "MCI"]


def test_las_tablas_heredan_la_seccion_y_su_rotulo(tmp_path: Path):
    docx = _docx()
    doc = docx.Document()
    doc.add_heading("Methods", level=1)
    doc.add_paragraph("Participants were recruited from three memory clinics.")
    doc.add_heading("Results", level=1)
    doc.add_paragraph("Table 1. Baseline characteristics by diagnostic group")
    primera = doc.add_table(rows=2, cols=2)
    primera.cell(0, 0).text = "Variable"
    primera.cell(0, 1).text = "Control"
    primera.cell(1, 0).text = "Edad"
    primera.cell(1, 1).text = "72.4"
    doc.add_heading("Discussion", level=1)
    doc.add_paragraph("These results should be read with caution.")
    segunda = doc.add_table(rows=1, cols=2)
    segunda.cell(0, 0).text = "Limitation"
    segunda.cell(0, 1).text = "Sample size"
    destino = tmp_path / "secciones.docx"
    doc.save(destino)

    chunks, _ = parse_generic(destino, destino.name)
    tablas = [c for c in chunks if c["chunk_type"] == "table"]

    assert [t["section"] for t in tablas] == ["Results", "Discussion"]
    assert [t["page"] for t in tablas] == [1, 2]
    assert tablas[0]["text"] == (
        "Results\nTable 1. Baseline characteristics by diagnostic group\n\n"
        "Variable | Control\nEdad | 72.4"
    )
    # Sin rótulo delante, solo la sección; un párrafo normal no cuenta como rótulo.
    assert tablas[1]["text"] == "Discussion\n\nLimitation | Sample size"


def test_una_tabla_larga_repite_la_cabecera_y_no_pierde_filas(tmp_path: Path):
    """Antes la tabla era un solo chunk recortado a 8000 caracteres: las filas
    del final desaparecían sin aviso."""
    docx = _docx()
    doc = docx.Document()
    doc.add_heading("Results", level=1)
    n = 300
    tabla = doc.add_table(rows=n + 1, cols=4)
    for j, valor in enumerate(["ID", "Grupo", "Basal", "Final"]):
        tabla.cell(0, j).text = valor
    for i in range(1, n + 1):
        for j, valor in enumerate([f"paciente{i:03d}", f"grupo{i % 3}", f"b{i}", f"f{i}"]):
            tabla.cell(i, j).text = valor
    destino = tmp_path / "larga.docx"
    doc.save(destino)

    chunks, _ = parse_generic(destino, destino.name)
    partes = [c for c in chunks if c["chunk_type"] == "table"]

    assert len(partes) > 1
    for k, parte in enumerate(partes, start=1):
        filas = _cuerpo(parte).splitlines()
        assert filas[0] == "ID | Grupo | Basal | Final"
        # Todas las partes citan la MISMA tabla y saben cuál son.
        assert parte["page"] == 1 and parte["section"] == "Results"
        assert parte["metadata"] == {"table_part": k, "table_parts": len(partes)}
    texto = _texto(partes)
    assert all(texto.count(f"paciente{i:03d} |") == 1 for i in range(1, n + 1))


def test_una_fila_de_titulo_no_sustituye_a_la_cabecera_en_cada_bloque(tmp_path: Path):
    """En Word la fila 0 de una tabla clínica suele ser el TÍTULO combinado a
    todo el ancho, y la cabecera real es la fila 1.

    Medido el 4 sep 2026 con una tabla de 200 filas: tomando ciegamente la fila
    0, las cinco partes salían encabezadas por "Table 1. Baseline
    characteristics" y SIN "ID | Grupo | Basal | Final", que es exactamente lo
    que la repetición de cabecera quiere evitar (en la parte 3, "b150" no
    significa nada sin el nombre de su columna).
    """
    docx = _docx()
    doc = docx.Document()
    doc.add_heading("Results", level=1)
    n = 200
    tabla = doc.add_table(rows=n + 2, cols=4)
    tabla.cell(0, 0).merge(tabla.cell(0, 3)).text = "Table 1. Baseline characteristics"
    for j, valor in enumerate(["ID", "Grupo", "Basal", "Final"]):
        tabla.cell(1, j).text = valor
    for i in range(1, n + 1):
        for j, valor in enumerate([f"paciente{i:03d}", f"grupo{i % 3}", f"b{i}", f"f{i}"]):
            tabla.cell(i + 1, j).text = valor
    destino = tmp_path / "titulo.docx"
    doc.save(destino)

    chunks, _ = parse_generic(destino, destino.name)
    partes = [c for c in chunks if c["chunk_type"] == "table"]

    assert len(partes) > 1
    for parte in partes:
        filas = _cuerpo(parte).splitlines()
        assert filas[0] == "Table 1. Baseline characteristics"
        assert filas[1] == "ID | Grupo | Basal | Final"
    # Y no se pierde ni se duplica ninguna fila de datos.
    texto = _texto(partes)
    assert all(texto.count(f"paciente{i:03d} |") == 1 for i in range(1, n + 1))


def test_una_cabecera_normal_sigue_siendo_una_sola_fila(tmp_path: Path):
    """Adversarial del criterio anterior: si la fila 0 nombra las columnas, la
    fila 1 son DATOS y no puede acabar dentro de la cabecera repetida (se
    duplicaría en cada bloque y desaparecería de su sitio)."""
    docx = _docx()
    doc = docx.Document()
    tabla = doc.add_table(rows=4, cols=3)
    for j, valor in enumerate(["Variable", "Control", "MCI"]):
        tabla.cell(0, j).text = valor
    datos = [["Edad", "72", "73"], ["MMSE", "28", "26"], ["Tau", "1.1", "2.4"]]
    for i, fila in enumerate(datos, start=1):
        for j, valor in enumerate(fila):
            tabla.cell(i, j).text = valor
    destino = tmp_path / "cabecera_normal.docx"
    doc.save(destino)

    chunks, _ = parse_generic(destino, destino.name)
    (tabla_chunk,) = [c for c in chunks if c["chunk_type"] == "table"]
    filas = _cuerpo(tabla_chunk).splitlines()

    assert filas == [
        "Variable | Control | MCI",
        "Edad | 72 | 73",
        "MMSE | 28 | 26",
        "Tau | 1.1 | 2.4",
    ]


def test_una_tabla_de_una_columna_no_pierde_su_primera_fila(tmp_path: Path):
    """Adversarial: con una sola columna la fila 0 tiene UNA celda efectiva y
    no es ningún título; si se tomara como cabecera de dos filas, la primera
    fila de datos quedaría absorbida."""
    docx = _docx()
    doc = docx.Document()
    tabla = doc.add_table(rows=3, cols=1)
    for i, valor in enumerate(["Criterios de exclusión", "Ictus previo", "Epilepsia"]):
        tabla.cell(i, 0).text = valor
    destino = tmp_path / "una_columna.docx"
    doc.save(destino)

    chunks, _ = parse_generic(destino, destino.name)
    (tabla_chunk,) = [c for c in chunks if c["chunk_type"] == "table"]

    assert _cuerpo(tabla_chunk).splitlines() == [
        "Criterios de exclusión",
        "Ictus previo",
        "Epilepsia",
    ]


def test_dos_tablas_consecutivas_no_comparten_el_rotulo(tmp_path: Path):
    """El rótulo describe UNA tabla: sin resetearlo al consumir una tabla, dos
    tablas seguidas heredaban las dos "Table 1. ..." y la segunda quedaba
    citada como la primera (medido el 4 sep 2026)."""
    docx = _docx()
    doc = docx.Document()
    doc.add_heading("Results", level=1)
    doc.add_paragraph("Table 1. Baseline characteristics")
    primera = doc.add_table(rows=1, cols=2)
    primera.cell(0, 0).text = "Edad"
    primera.cell(0, 1).text = "72.4"
    segunda = doc.add_table(rows=1, cols=2)
    segunda.cell(0, 0).text = "MMSE"
    segunda.cell(0, 1).text = "28.1"
    destino = tmp_path / "consecutivas.docx"
    doc.save(destino)

    chunks, _ = parse_generic(destino, destino.name)
    tablas = [c for c in chunks if c["chunk_type"] == "table"]

    assert tablas[0]["text"] == "Results\nTable 1. Baseline characteristics\n\nEdad | 72.4"
    assert tablas[1]["text"] == "Results\n\nMMSE | 28.1"


def test_los_parrafos_de_word_se_solapan_y_llevan_la_seccion(tmp_path: Path):
    docx = _docx()
    doc = docx.Document()
    doc.add_heading("Métodos", level=1)
    for i in range(40):
        doc.add_paragraph(
            f"Paso {i:03d} del reclutamiento, descrito con el detalle suficiente "
            f"para que el parrafo ocupe unas veinte palabras completas."
        )
    doc.add_heading("Resultados", level=1)
    doc.add_paragraph("La media de amiloide fue 542 pg/mL en el grupo afectado.")
    destino = tmp_path / "solape.docx"
    doc.save(destino)

    chunks, _ = parse_generic(destino, destino.name)
    metodos = _de_seccion(chunks, "Métodos")

    assert len(metodos) >= 2
    # El primero arranca por el encabezado indexado como bloque, una sola vez.
    assert metodos[0]["text"].startswith("Métodos\n\nPaso 000")
    assert metodos[0]["text"].count("Métodos") == 1
    # Los siguientes llevan la sección delante y la cola del anterior detrás.
    for anterior, siguiente in zip(metodos, metodos[1:]):
        assert siguiente["text"].startswith("Métodos\n\n")
        cola = _solape(anterior, siguiente)
        assert cola
        assert _OVERLAP_TOKENS // 2 <= sum(_est_tokens(p) for p in cola) <= 2 * _OVERLAP_TOKENS
    (resultados,) = _de_seccion(chunks, "Resultados")
    assert "reclutamiento" not in resultados["text"]
    assert resultados["text"].startswith("Resultados\n\nLa media")
